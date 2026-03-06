#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档格式转换器
提供 Excel、Word、PPT、PDF 等格式之间的转换
"""

import pandas as pd
import json
import csv
from pathlib import Path


class DocumentConverter:
    """文档格式转换器"""
    
    # ==================== Excel 转换 ====================
    
    @staticmethod
    def excel_to_csv(excel_path, csv_path, sheet_name=0):
        """Excel 转 CSV"""
        df = pd.read_excel(excel_path, sheet_name=sheet_name)
        df.to_csv(csv_path, index=False, encoding='utf-8-sig')
        print(f"✅ 已转换: {excel_path} → {csv_path}")
        return csv_path
    
    @staticmethod
    def excel_to_json(excel_path, json_path, sheet_name=0):
        """Excel 转 JSON"""
        df = pd.read_excel(excel_path, sheet_name=sheet_name)
        data = df.to_dict(orient='records')
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"✅ 已转换: {excel_path} → {json_path}")
        return json_path
    
    @staticmethod
    def excel_to_markdown(excel_path, md_path, sheet_name=0):
        """Excel 转 Markdown"""
        df = pd.read_excel(excel_path, sheet_name=sheet_name)
        
        # 生成 Markdown 表格
        lines = []
        lines.append(f"# {Path(excel_path).stem}")
        lines.append("")
        
        # 表头
        headers = list(df.columns)
        line = "| " + " | ".join(str(h) for h in headers) + " |"
        lines.append(line)
        lines.append("|" + "|".join(["---"] * len(headers)) + "|")
        
        # 数据行
        for _, row in df.iterrows():
            line = "| " + " | ".join(str(row[h]) for h in headers) + " |"
            lines.append(line)
        
        md_content = '\n'.join(lines)
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        print(f"✅ 已转换: {excel_path} → {md_path}")
        return md_path
    
    @staticmethod
    def csv_to_excel(csv_path, excel_path, sheet_name='Sheet1'):
        """CSV 转 Excel"""
        df = pd.read_csv(csv_path, encoding='utf-8-sig')
        df.to_excel(excel_path, sheet_name=sheet_name, index=False)
        print(f"✅ 已转换: {csv_path} → {excel_path}")
        return excel_path
    
    @staticmethod
    def json_to_excel(json_path, excel_path, sheet_name='Sheet1'):
        """JSON 转 Excel"""
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        df = pd.DataFrame(data)
        df.to_excel(excel_path, sheet_name=sheet_name, index=False)
        print(f"✅ 已转换: {json_path} → {excel_path}")
        return excel_path
    
    # ==================== Word 转换 ====================
    
    @staticmethod
    def word_to_text(word_path, text_path):
        """Word 转纯文本"""
        from docx import Document
        doc = Document(word_path)
        
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        
        text_content = '\n'.join(text_parts)
        with open(text_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        
        print(f"✅ 已转换: {word_path} → {text_path}")
        return text_path
    
    @staticmethod
    def extract_word_tables(word_path):
        """从 Word 提取表格数据"""
        from docx import Document
        doc = Document(word_path)
        
        tables_data = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            tables_data.append(table_data)
        
        print(f"✅ 从 {word_path} 提取了 {len(tables_data)} 个表格")
        return tables_data
    
    # ==================== 多文件批量转换 ====================
    
    @staticmethod
    def batch_excel_to_csv(excel_dir, csv_dir):
        """批量 Excel 转 CSV"""
        Path(csv_dir).mkdir(parents=True, exist_ok=True)
        
        converted = []
        for excel_file in Path(excel_dir).glob('*.xlsx'):
            csv_file = Path(csv_dir) / (excel_file.stem + '.csv')
            DocumentConverter.excel_to_csv(str(excel_file), str(csv_file))
            converted.append(str(csv_file))
        
        print(f"✅ 批量转换完成: {len(converted)} 个文件")
        return converted
    
    @staticmethod
    def batch_csv_to_excel(csv_dir, excel_path):
        """批量 CSV 合并为 Excel（多 sheet）"""
        with pd.ExcelWriter(excel_path) as writer:
            for csv_file in Path(csv_dir).glob('*.csv'):
                df = pd.read_csv(csv_file, encoding='utf-8-sig')
                sheet_name = csv_file.stem[:31]  # Excel sheet 名称限制31字符
                df.to_excel(writer, sheet_name=sheet_name, index=False)
        
        print(f"✅ 已合并到: {excel_path}")
        return excel_path
    
    # ==================== 数据导出 ====================
    
    @staticmethod
    def export_to_format(data, output_path, format='csv'):
        """通用数据导出"""
        if format == 'csv':
            if isinstance(data, list):
                df = pd.DataFrame(data)
            else:
                df = data
            df.to_csv(output_path, index=False, encoding='utf-8-sig')
        
        elif format == 'excel':
            if isinstance(data, list):
                df = pd.DataFrame(data)
            else:
                df = data
            df.to_excel(output_path, index=False)
        
        elif format == 'json':
            if isinstance(data, pd.DataFrame):
                data = data.to_dict(orient='records')
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        
        elif format == 'markdown':
            if isinstance(data, pd.DataFrame):
                # 生成 Markdown 表格
                lines = []
                headers = list(data.columns)
                line = "| " + " | ".join(str(h) for h in headers) + " |"
                lines.append(line)
                lines.append("|" + "|".join(["---"] * len(headers)) + "|")
                
                for _, row in data.iterrows():
                    line = "| " + " | ".join(str(row[h]) for h in headers) + " |"
                    lines.append(line)
                
                content = '\n'.join(lines)
            else:
                content = str(data)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
        
        print(f"✅ 已导出: {output_path}")
        return output_path
    
    # ==================== 数据导入 ====================
    
    @staticmethod
    def import_from_format(input_path):
        """通用数据导入"""
        suffix = Path(input_path).suffix.lower()
        
        if suffix == '.csv':
            return pd.read_csv(input_path, encoding='utf-8-sig')
        
        elif suffix in ['.xlsx', '.xls']:
            return pd.read_excel(input_path)
        
        elif suffix == '.json':
            with open(input_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        
        elif suffix == '.txt':
            with open(input_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        else:
            raise ValueError(f"不支持的格式: {suffix}")
    
    # ==================== 报告生成 ====================
    
    @staticmethod
    def generate_summary_report(data_files, report_path):
        """生成汇总报告"""
        report = []
        report.append("# 数据汇总报告")
        report.append(f"生成时间: 2024-01-20")
        report.append("")
        
        for file_path in data_files:
            file = Path(file_path)
            report.append(f"## {file.name}")
            
            try:
                df = DocumentConverter.import_from_format(file_path)
                report.append(f"- 行数: {len(df)}")
                report.append(f"- 列数: {len(df.columns)}")
                report.append(f"- 列名: {', '.join(df.columns)}")
            except Exception as e:
                report.append(f"- 读取失败: {e}")
            
            report.append("")
        
        content = '\n'.join(report)
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        print(f"✅ 汇总报告已生成: {report_path}")
        return report_path


# ==================== 便捷函数 ====================

def excel2csv(excel_path, csv_path):
    """快速 Excel 转 CSV"""
    return DocumentConverter.excel_to_csv(excel_path, csv_path)


def excel2json(excel_path, json_path):
    """快速 Excel 转 JSON"""
    return DocumentConverter.excel_to_json(excel_path, json_path)


def csv2excel(csv_path, excel_path):
    """快速 CSV 转 Excel"""
    return DocumentConverter.csv_to_excel(csv_path, excel_path)


def json2excel(json_path, excel_path):
    """快速 JSON 转 Excel"""
    return DocumentConverter.json_to_excel(json_path, excel_path)


def batch_convert(excel_dir, output_dir, format='csv'):
    """批量转换"""
    if format == 'csv':
        return DocumentConverter.batch_excel_to_csv(excel_dir, output_dir)


if __name__ == "__main__":
    print("🧪 测试文档转换器")
    
    import os
    import tempfile
    
    # 创建测试数据
    with tempfile.TemporaryDirectory() as tmpdir:
        # 1. 创建测试 Excel
        excel_file = os.path.join(tmpdir, "test.xlsx")
        df = pd.DataFrame({
            '姓名': ['张三', '李四', '王五'],
            '部门': ['技术部', '产品部', '设计部'],
            '分数': [85, 90, 78]
        })
        df.to_excel(excel_file, index=False)
        
        # 2. 测试 Excel 转 CSV
        csv_file = os.path.join(tmpdir, "test.csv")
        DocumentConverter.excel_to_csv(excel_file, csv_file)
        
        # 3. 测试 Excel 转 JSON
        json_file = os.path.join(tmpdir, "test.json")
        DocumentConverter.excel_to_json(excel_file, json_file)
        
        # 4. 测试 Excel 转 Markdown
        md_file = os.path.join(tmpdir, "test.md")
        DocumentConverter.excel_to_markdown(excel_file, md_file)
        
        # 5. 测试 CSV 转 Excel
        csv2excel_file = os.path.join(tmpdir, "csv2excel.xlsx")
        DocumentConverter.csv_to_excel(csv_file, csv2excel_file)
        
        # 6. 测试 JSON 转 Excel
        json2excel_file = os.path.join(tmpdir, "json2excel.xlsx")
        DocumentConverter.json_to_excel(json_file, json2excel_file)
        
        print("\n✅ 所有转换测试完成！")
        print(f"📁 测试文件位于: {tmpdir}")
