#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 文档处理器
提供 Word 文档创建、编辑、格式化、报告生成等功能
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from datetime import datetime
from pathlib import Path


class WordProcessor:
    """Word 文档处理类"""
    
    def __init__(self):
        self.doc = None
        self.filepath = None
    
    # ==================== 基础操作 ====================
    
    def create(self):
        """创建新文档"""
        self.doc = Document()
        return self
    
    def load(self, filepath):
        """加载现有文档"""
        self.doc = Document(filepath)
        self.filepath = filepath
        return self
    
    def save(self, filepath=None):
        """保存文档"""
        save_path = filepath or self.filepath
        if save_path:
            self.doc.save(save_path)
            print(f"✅ 已保存: {save_path}")
        else:
            print("⚠️ 未指定保存路径")
        return save_path
    
    # ==================== 文本添加 ====================
    
    def add_heading(self, text, level=0):
        """添加标题"""
        self.doc.add_heading(text, level)
        return self
    
    def add_paragraph(self, text=None, style=None):
        """添加段落"""
        para = self.doc.add_paragraph(text, style)
        return para
    
    def add_bullet(self, text, level=0):
        """添加项目符号"""
        para = self.doc.add_paragraph(text, style='List Bullet')
        if level > 0:
            para.level = level
        return para
    
    def add_numbered(self, text, level=0):
        """添加编号"""
        para = self.doc.add_paragraph(text, style='List Number')
        if level > 0:
            para.level = level
        return para
    
    # ==================== 格式化 ====================
    
    def set_default_font(self, font_name='微软雅黑', font_size=12):
        """设置默认字体"""
        style = self.doc.styles['Normal']
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        return self
    
    def format_text(self, text, bold=False, italic=False, size=None, color=None):
        """格式化文本"""
        run = self.doc.add_paragraph().add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if color:
            if isinstance(color, tuple):
                run.font.color.rgb = RGBColor(*color)
        return run
    
    # ==================== 表格操作 ====================
    
    def add_table(self, rows, cols, data=None, style='Table Grid'):
        """创建表格"""
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = style
        
        if data:
            for row_idx, row_data in enumerate(data):
                for col_idx, value in enumerate(row_data):
                    cell = table.cell(row_idx, col_idx)
                    cell.text = str(value)
                    # 设置表格自动换行
                    cell.text_frame.paragraph_format.word_wrap = True
        
        return table
    
    def add_project_table(self, project_info):
        """创建项目信息表格"""
        data = [
            ['项目属性', '内容'],
            ['项目名称', project_info.get('name', '')],
            ['项目经理', project_info.get('manager', '')],
            ['开始日期', project_info.get('start_date', '')],
            ['结束日期', project_info.get('end_date', '')],
            ['预算', project_info.get('budget', '')],
            ['状态', project_info.get('status', '')]
        ]
        return self.add_table(7, 2, data)
    
    # ==================== 图片操作 ====================
    
    def add_image(self, image_path, width=None, height=None):
        """添加图片"""
        if width:
            self.doc.add_picture(image_path, width=Inches(width))
        elif height:
            self.doc.add_picture(image_path, height=Inches(height))
        else:
            self.doc.add_picture(image_path)
        return self
    
    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()
        return self
    
    # ==================== 模板生成 ====================
    
    @staticmethod
    def create_project_report(filepath, project_info, progress_data):
        """创建项目报告"""
        doc = Document()
        
        # 标题
        doc.add_heading(f"{project_info['name']} - 项目报告", 0)
        
        # 基本信息
        doc.add_heading('项目概述', 1)
        doc.add_paragraph(f"项目经理：{project_info['manager']}")
        doc.add_paragraph(f"报告周期：{project_info['period']}")
        doc.add_paragraph(f"当前状态：{project_info['status']}")
        doc.add_paragraph(f"整体进度：{project_info['progress']}%")
        
        doc.add_page_break()
        
        # 进度摘要
        doc.add_heading('进度摘要', 1)
        
        # 进度表格
        if progress_data:
            table = doc.add_table(rows=len(progress_data)+1, cols=4)
            table.style = 'Table Grid'
            
            # 表头
            headers = ['任务', '负责人', '进度', '状态']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            
            # 数据
            for row_idx, task in enumerate(progress_data):
                row = table.rows[row_idx + 1].cells
                row[0].text = task['name']
                row[1].text = task['owner']
                row[2].text = f"{task['progress']}%"
                row[3].text = task['status']
        
        doc.add_page_break()
        
        # 问题与风险
        doc.add_heading('问题与风险', 1)
        risks = project_info.get('risks', [])
        for risk in risks:
            doc.add_paragraph(f"⚠️ {risk}", style='List Bullet')
        
        # 下一步计划
        doc.add_heading('下一步计划', 1)
        plans = project_info.get('plans', [])
        for i, plan in enumerate(plans):
            doc.add_paragraph(f"{i+1}. {plan}", style='List Number')
        
        # 保存
        doc.save(filepath)
        print(f"✅ 项目报告已创建: {filepath}")
        return filepath
    
    @staticmethod
    def create_weekly_report(filepath, week_info):
        """创建周报"""
        doc = Document()
        
        # 标题
        doc.add_heading(f"{week_info['project']} - 周报", 0)
        doc.add_paragraph(f"日期范围: {week_info['date_range']}")
        doc.add_paragraph(f"周次: {week_info['week']}")
        
        # 本周完成
        doc.add_heading('本周完成', 1)
        for task in week_info.get('completed', []):
            doc.add_paragraph(f"✅ {task}", style='List Bullet')
        
        # 本周进行
        doc.add_heading('进行中', 1)
        for task in week_info.get('in_progress', []):
            doc.add_paragraph(f"🔄 {task}", style='List Bullet')
        
        # 问题
        doc.add_heading('遇到问题', 1)
        for issue in week_info.get('issues', []):
            doc.add_paragraph(f"❓ {issue}", style='List Bullet')
        
        # 下周计划
        doc.add_heading('下周计划', 1)
        for plan in week_info.get('next_week', []):
            doc.add_paragraph(f"📋 {plan}", style='List Bullet')
        
        doc.save(filepath)
        print(f"✅ 周报已创建: {filepath}")
        return filepath
    
    @staticmethod
    def create_prd(filepath, prd_info):
        """创建产品需求文档"""
        doc = Document()
        
        # 封面
        doc.add_heading(prd_info['product_name'], 0)
        doc.add_paragraph(f"版本: {prd_info['version']}")
        doc.add_paragraph(f"作者: {prd_info['author']}")
        doc.add_paragraph(f"日期: {prd_info['date']}")
        
        doc.add_page_break()
        
        # 目录
        doc.add_heading('目录', 1)
        sections = ['概述', '用户故事', '功能需求', '非功能需求', '交互设计', '验收标准']
        for i, section in enumerate(sections):
            doc.add_paragraph(f"{i+1}. {section}", style='List Number')
        
        doc.add_page_break()
        
        # 各章节
        for section in sections:
            doc.add_heading(f"{section}", 1)
            doc.add_paragraph("...")
        
        doc.save(filepath)
        print(f"✅ PRD 已创建: {filepath}")
        return filepath
    
    @staticmethod
    def create_meeting_minutes(filepath, meeting_info):
        """创建会议纪要"""
        doc = Document()
        
        doc.add_heading(f"{meeting_info['title']}", 0)
        
        doc.add_paragraph(f"日期: {meeting_info['date']}")
        doc.add_paragraph(f"主持人: {meeting_info['host']}")
        doc.add_paragraph(f"参会人: {meeting_info['attendees']}")
        
        doc.add_heading('议程', 1)
        for agenda in meeting_info.get('agenda', []):
            doc.add_paragraph(agenda, style='List Number')
        
        doc.add_heading('讨论内容', 1)
        for topic in meeting_info.get('discussions', []):
            doc.add_heading(topic['topic'], 2)
            doc.add_paragraph(topic['content'])
        
        doc.add_heading('决议', 1)
        for decision in meeting_info.get('decisions', []):
            doc.add_paragraph(f"✅ {decision}", style='List Bullet')
        
        doc.add_heading('待办事项', 1)
        for todo in meeting_info.get('todos', []):
            doc.add_paragraph(f"📋 {todo}", style='List Bullet')
        
        doc.save(filepath)
        print(f"✅ 会议纪要已创建: {filepath}")
        return filepath
    
    @staticmethod
    def create_template_report(filepath, title, sections):
        """创建通用模板报告"""
        doc = Document()
        
        doc.add_heading(title, 0)
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        for section in sections:
            doc.add_heading(section['title'], 1)
            
            if section['type'] == 'paragraph':
                doc.add_paragraph(section['content'])
            elif section['type'] == 'list':
                for item in section['items']:
                    doc.add_paragraph(item, style='List Bullet')
            elif section['type'] == 'numbered':
                for i, item in enumerate(section['items']):
                    doc.add_paragraph(f"{i+1}. {item}", style='List Number')
            elif section['type'] == 'table':
                doc.add_table(section['rows'], section['cols'], section['data'])
        
        doc.save(filepath)
        print(f"✅ 报告已创建: {filepath}")
        return filepath


# ==================== 便捷函数 ====================

def quick_create(title):
    """快速创建文档"""
    return WordProcessor().create().add_heading(title, 0)


def create_project_status_report(filepath, project_name, tasks):
    """快速创建项目状态报告"""
    info = {
        'name': project_name,
        'manager': '项目经理',
        'period': '本周期',
        'status': '进行中',
        'progress': 50,
        'risks': ['风险1', '风险2'],
        'plans': ['计划1', '计划2']
    }
    return WordProcessor.create_project_report(filepath, info, tasks)


if __name__ == "__main__":
    print("🧪 测试 Word 处理器")
    
    # 测试创建文档
    print("\n1. 创建测试文档...")
    test_file = "/tmp/test_document.docx"
    doc = WordProcessor()
    doc.create()
    doc.set_default_font()
    doc.add_heading("测试文档", 0)
    doc.add_paragraph("这是一个测试段落")
    doc.add_bullet("要点1")
    doc.add_bullet("要点2")
    doc.add_table(3, 3, [['A', 'B', 'C'], ['1', '2', '3'], ['X', 'Y', 'Z']])
    doc.save(test_file)
    
    print("\n2. 测试创建项目报告...")
    report_file = "/tmp/test_report.docx"
    WordProcessor.create_project_report(
        report_file,
        {
            'name': '测试项目',
            'manager': '张三',
            'period': '2024年1月',
            'status': '进行中',
            'progress': 60,
            'risks': ['技术风险', '人员风险'],
            'plans': ['完成开发', '开始测试']
        },
        [
            {'name': '需求分析', 'owner': '李四', 'progress': 100, 'status': '已完成'},
            {'name': '系统设计', 'owner': '王五', 'progress': 80, 'status': '进行中'},
            {'name': '开发实施', 'owner': '赵六', 'progress': 30, 'status': '进行中'}
        ]
    )
    
    print("\n3. 测试完成！")
